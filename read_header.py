import docx

# Ouvrir le document
doc_path = r"d:\NGAPEP01\OneDrive - Castel Afrique\Documents\Projet SABC\Projet 6.0\Hyper-Framework\Contrôles SSI  - S42.docx"
doc = docx.Document(doc_path)

print("="*80)
print("EN-TÊTE DU DOCUMENT")
print("="*80)

# Parcourir toutes les sections
for section_idx, section in enumerate(doc.sections):
    print(f"\n\n### SECTION {section_idx + 1} ###")
    header = section.header
    
    print(f"\nNombre de paragraphes dans l'en-tête: {len(header.paragraphs)}")
    
    # Lire tous les paragraphes
    for i, para in enumerate(header.paragraphs):
        if para.text.strip():  # Seulement si non vide
            print(f"\n--- Paragraphe {i+1} ---")
            print(f"Texte: {para.text}")
            print(f"Style: {para.style.name}")
            print(f"Alignement: {para.alignment}")
            
            # Détails des runs
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    print(f"  Run {run_idx+1}: '{run.text}'")
                    print(f"    Gras: {run.bold}, Italique: {run.italic}")
                    if run.font.size:
                        print(f"    Taille: {run.font.size}")
    
    # Lire les tables dans l'en-tête
    print(f"\n\nNombre de tables dans l'en-tête: {len(header.tables)}")
    
    for table_idx, table in enumerate(header.tables):
        print(f"\n--- Table {table_idx + 1} ---")
        print(f"Dimensions: {len(table.rows)} lignes x {len(table.columns)} colonnes")
        
        for row_idx, row in enumerate(table.rows):
            print(f"\n  Ligne {row_idx + 1}:")
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"    Colonne {cell_idx + 1}: {cell_text}")

print("\n" + "="*80)
