#---> FICHIER MODIFIÉ : hyper_framework_server/services/report_service.py

from datetime import datetime
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')  # Mode non-interactif
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import os
import tempfile
import re
from ..config import DEFAULT_ASSETS_DIR

class ReportGenerator:
    def _create_conclusion_hexagon(self, compliance_rate):
        """
        Crée un hexagone coloré selon le taux de conformité.
        
        Args:
            compliance_rate: Taux de conformité en pourcentage (0-100)
            
        Returns:
            Chemin vers l'image temporaire de l'hexagone
        """
        # Déterminer la couleur selon le taux
        if compliance_rate >= 95:
            color = '#4CAF50'  # Vert
        elif compliance_rate >= 50:
            color = '#FFC107'  # Jaune/Orange
        else:
            color = '#F44336'  # Rouge
        
        # Créer la figure
        fig, ax = plt.subplots(figsize=(0.5, 0.5), dpi=150)
        ax.set_xlim(-1.2, 1.2)
        ax.set_ylim(-1.2, 1.2)
        ax.axis('off')
        
        # Créer l'hexagone
        hexagon = mpatches.RegularPolygon(
            (0, 0),  # Centre
            6,  # Nombre de côtés
            radius=1,
            facecolor=color,
            edgecolor='black',
            linewidth=1
        )
        ax.add_patch(hexagon)
        
        # Sauvegarder dans un fichier temporaire
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        plt.savefig(temp_file.name, format='png', dpi=150, bbox_inches='tight', pad_inches=0)
        plt.close()
        
        return temp_file.name
    def _add_header_with_logo_and_table(self, document, control_name, control_code, analysis_results, control_metadata=None, execution_date=None):
        """
        Ajoute un en-tête avec une structure en tableau unique de 4 lignes x 7 colonnes
        
        Args:
            control_metadata: Dictionnaire contenant les métadonnées du contrôle
            execution_date: Date d'exécution au format YYYYMMDD-HHMMSS
        """
        section = document.sections[0]
        header = section.header
        
        # Nettoyer l'en-tête existant
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Créer UN SEUL tableau de 4 lignes x 7 colonnes
        header_table = header.add_table(rows=4, cols=7, width=Inches(7.0))
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # --- LIGNE 1 : Titre du contrôle (colonnes 2-5) et Code (colonnes 6-7) ---
        # Colonne 1 : Logo
        logo_cell = header_table.rows[0].cells[0]
        logo_para = logo_cell.paragraphs[0]
        logo_para.paragraph_format.space_before = Pt(4)
        logo_para.paragraph_format.space_after = Pt(4)
        logo_para.paragraph_format.line_spacing = 1.15
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le logo avec les dimensions spécifiées (2,81 cm x 1,42 cm)
        logo_path = DEFAULT_ASSETS_DIR / 'images' / 'logo_default.png'
        if os.path.exists(logo_path):
            run = logo_para.add_run()
            # 2,81 cm = 1.106 pouces, 1,42 cm = 0.559 pouces
            run.add_picture(str(logo_path), width=Inches(1.106), height=Inches(0.559))
        
        # Centrer verticalement
        tc_logo = logo_cell._element
        tcPr_logo = tc_logo.get_or_add_tcPr()
        tcVAlign_logo = OxmlElement('w:vAlign')
        tcVAlign_logo.set(qn('w:val'), 'center')
        tcPr_logo.append(tcVAlign_logo)
        
        # Fusionner colonnes 2-5 pour le titre
        title_cell = header_table.rows[0].cells[1]
        for i in range(2, 5):
            title_cell.merge(header_table.rows[0].cells[i])
        
        title_para = title_cell.paragraphs[0]
        title_para.paragraph_format.space_before = Pt(4)
        title_para.paragraph_format.space_after = Pt(4)
        title_para.paragraph_format.line_spacing = 1.15
        title_run = title_para.add_run(control_name)
        title_run.font.name = 'Tahoma'
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc_title = title_cell._element
        tcPr_title = tc_title.get_or_add_tcPr()
        tcVAlign_title = OxmlElement('w:vAlign')
        tcVAlign_title.set(qn('w:val'), 'center')
        tcPr_title.append(tcVAlign_title)
        
        # Fusionner colonnes 6-7 pour le code
        # Construire le code avec la date d'exécution
        if control_metadata and execution_date:
            control_code_prefix = control_metadata.get('control_code_prefix', control_code)
            # Extraire la date du timestamp (format: YYYYMMDD-HHMMSS)
            date_part = execution_date.split('-')[0]  # YYYYMMDD
            # Convertir en format YYYY_MM_DD
            formatted_date = f"{date_part[:4]}_{date_part[4:6]}_{date_part[6:8]}"
            full_control_code = f"{control_code_prefix}_{formatted_date}"
        else:
            full_control_code = control_code
        
        code_cell = header_table.rows[0].cells[5]
        code_cell.merge(header_table.rows[0].cells[6])
        
        code_para = code_cell.paragraphs[0]
        code_para.paragraph_format.space_before = Pt(4)
        code_para.paragraph_format.space_after = Pt(4)
        code_para.paragraph_format.line_spacing = 1.15
        code_run = code_para.add_run(full_control_code)
        code_run.font.name = 'Tahoma'
        code_run.font.size = Pt(6)
        code_run.font.bold = True
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc_code = code_cell._element
        tcPr_code = tc_code.get_or_add_tcPr()
        tcVAlign_code = OxmlElement('w:vAlign')
        tcVAlign_code.set(qn('w:val'), 'center')
        tcPr_code.append(tcVAlign_code)
        
        # --- LIGNE 2 : En-têtes des colonnes ---
        headers = [
            'Application concernée',
            'Couche concernée',
            'Référence du risque',
            'Nom du risque',
            'Conclusion'
        ]
        
        # Remplir les 5 premières colonnes
        for i, header_text in enumerate(headers):
            cell = header_table.rows[1].cells[i]
            cell_para = cell.paragraphs[0]
            cell_para.paragraph_format.space_before = Pt(4)
            cell_para.paragraph_format.space_after = Pt(4)
            cell_para.paragraph_format.line_spacing = 1.15
            cell_run = cell_para.add_run(header_text)
            cell_run.font.name = 'Tahoma'
            cell_run.font.size = Pt(6)
            cell_run.font.bold = True
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Centrer verticalement
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), 'center')
            tcPr.append(tcVAlign)
        
        # Fusionner les colonnes 6 et 7 pour "Nom du contrôle"
        control_name_cell = header_table.rows[1].cells[5]
        control_name_cell.merge(header_table.rows[1].cells[6])
        
        cell_para = control_name_cell.paragraphs[0]
        cell_para.paragraph_format.space_before = Pt(4)
        cell_para.paragraph_format.space_after = Pt(4)
        cell_para.paragraph_format.line_spacing = 1.15
        cell_run = cell_para.add_run('Nom du contrôle')
        cell_run.font.name = 'Tahoma'
        cell_run.font.size = Pt(6)
        cell_run.font.bold = True
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc = control_name_cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # --- LIGNE 3 : Données du contrôle ---
        # Utiliser les métadonnées du contrôle si disponibles
        if control_metadata:
            application = control_metadata.get('application', 'N/A')
            layer = control_metadata.get('layer', 'N/A')
            risk_reference = control_metadata.get('risk_reference', 'N/A')
            risk_name = control_metadata.get('risk_name', 'N/A')
            control_name_meta = control_metadata.get('control_name', control_name)
        else:
            application = 'N/A'
            layer = 'N/A'
            risk_reference = 'N/A'
            risk_name = 'N/A'
            control_name_meta = control_name
        
        # Calculer le taux de conformité pour la conclusion (hexagone)
        compliance_rate = 0
        if analysis_results and len(analysis_results) > 0:
            first_section = analysis_results[0]
            summary_stats = first_section.get('summary_stats', {})
            
            for key, value in summary_stats.items():
                if 'taux' in key.lower() or 'conformité' in key.lower() or 'conformite' in key.lower():
                    if isinstance(value, str):
                        match = re.search(r'(\d+\.?\d*)', value)
                        if match:
                            compliance_rate = float(match.group(1))
                            break
                    elif isinstance(value, (int, float)):
                        compliance_rate = float(value)
                        break
        
        # Colonne 1 : Application
        cell = header_table.rows[2].cells[0]
        cell.text = application
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = 'Tahoma'
            run.font.size = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # Colonne 2 : Couche
        cell = header_table.rows[2].cells[1]
        cell.text = layer
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = 'Tahoma'
            run.font.size = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # Colonne 3 : Référence du risque
        cell = header_table.rows[2].cells[2]
        cell.text = risk_reference
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = 'Tahoma'
            run.font.size = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # Colonne 4 : Nom du risque
        cell = header_table.rows[2].cells[3]
        cell.text = risk_name
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = 'Tahoma'
            run.font.size = Pt(5.5)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # Colonne 5 : Conclusion (vide pour le modèle S42)
        cell = header_table.rows[2].cells[4]
        cell.text = ''
        
        # Fusionner colonnes 6-7 pour Nom du contrôle
        control_name_data_cell = header_table.rows[2].cells[5]
        control_name_data_cell.merge(header_table.rows[2].cells[6])
        
        control_name_data_cell.text = control_name_meta
        para = control_name_data_cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.15
        for run in para.runs:
            run.font.name = 'Tahoma'
            run.font.size = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = control_name_data_cell._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # --- LIGNE 4 : Destinataire, Ref Description, PS3 ---
        # Fusionner colonnes 1-3 pour Destinataire
        dest_cell = header_table.rows[3].cells[0]
        dest_cell.merge(header_table.rows[3].cells[1])
        dest_cell.merge(header_table.rows[3].cells[2])
        
        dest_para = dest_cell.paragraphs[0]
        dest_para.paragraph_format.space_before = Pt(4)
        dest_para.paragraph_format.space_after = Pt(4)
        dest_para.paragraph_format.line_spacing = 1.15
        dest_run1 = dest_para.add_run('Destinataire : ')
        dest_run1.font.name = 'Tahoma'
        dest_run1.font.size = Pt(6)
        dest_run1.font.bold = True
        dest_run2 = dest_para.add_run('Tout le personnel de la direction des systèmes d\'informations')
        dest_run2.font.name = 'Tahoma'
        dest_run2.font.size = Pt(6)
        
        tc_dest = dest_cell._element
        tcPr_dest = tc_dest.get_or_add_tcPr()
        tcVAlign_dest = OxmlElement('w:vAlign')
        tcVAlign_dest.set(qn('w:val'), 'center')
        tcPr_dest.append(tcVAlign_dest)
        
        # Fusionner colonnes 4-5 pour Ref Description
        ref_cell = header_table.rows[3].cells[3]
        ref_cell.merge(header_table.rows[3].cells[4])
        
        ref_para = ref_cell.paragraphs[0]
        ref_para.paragraph_format.space_before = Pt(4)
        ref_para.paragraph_format.space_after = Pt(4)
        ref_para.paragraph_format.line_spacing = 1.15
        ref_run1 = ref_para.add_run('Ref Descirption : ')  # Garder la faute comme dans le modèle
        ref_run1.font.name = 'Tahoma'
        ref_run1.font.size = Pt(6)
        ref_run1.font.bold = True

        ref_description = 'N/A'
        if control_metadata:
            ref_description = control_metadata.get('ref_description', 'N/A')

        ref_run2 = ref_para.add_run(ref_description)
        ref_run2.font.name = 'Tahoma'
        ref_run2.font.size = Pt(6)
        
        tc_ref = ref_cell._element
        tcPr_ref = tc_ref.get_or_add_tcPr()
        tcVAlign_ref = OxmlElement('w:vAlign')
        tcVAlign_ref.set(qn('w:val'), 'center')
        tcPr_ref.append(tcVAlign_ref)
        
        # Fusionner colonnes 6-7 pour PS3
        ps_cell = header_table.rows[3].cells[5]
        ps_cell.merge(header_table.rows[3].cells[6])
        
        ps_para = ps_cell.paragraphs[0]
        ps_para.paragraph_format.space_before = Pt(4)
        ps_para.paragraph_format.space_after = Pt(4)
        ps_para.paragraph_format.line_spacing = 1.15
        ps_run = ps_para.add_run('PS3_BC_SSI_FEN_3')
        ps_run.font.name = 'Tahoma'
        ps_run.font.size = Pt(6)
        ps_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        tc_ps = ps_cell._element
        tcPr_ps = tc_ps.get_or_add_tcPr()
        tcVAlign_ps = OxmlElement('w:vAlign')
        tcVAlign_ps.set(qn('w:val'), 'center')
        tcPr_ps.append(tcVAlign_ps)
        
        # --- Configurer les bordures et dimensions du tableau ---
        tbl = header_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Centrer le tableau
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)
        
        # Définir la largeur totale du tableau 
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '10195')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Layout fixe pour empêcher l'auto-ajustement
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # --- Définiton des largeurs des colonnes via la grille ET les cellules ---
        # 1 cm = 567 twips

        # Définition de la grille
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            gridCols = tblGrid.findall(qn('w:gridCol'))
            if len(gridCols) >= 7:
                gridCols[0].set(qn('w:w'), '2138')  # 3,77 cm
                gridCols[1].set(qn('w:w'), '1418')  # 2,5 cm
                gridCols[2].set(qn('w:w'), '1530')  # 2,7 cm
                gridCols[3].set(qn('w:w'), '1418')  # 2,5 cm
                gridCols[4].set(qn('w:w'), '1417')  # 2,5 cm
                gridCols[5].set(qn('w:w'), '1137')  # 2,005 cm
                gridCols[6].set(qn('w:w'), '1137')  # 2,005 cm
        
        # Respect des dimensions sur Word
        col_widths = [2138, 1418, 1530, 1418, 1417, 1137, 1137]
        for idx, cell in enumerate(header_table.rows[1].cells[:7]):
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[idx]))
            tcW.set(qn('w:type'), 'dxa')
            # Supprimer l'ancien tcW s'il existe
            old_tcW = tcPr.find(qn('w:tcW'))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcPr.append(tcW)
    
    def _add_signature_table(self, document):
        """
        Ajoute le tableau de signatures à la fin du rapport.
        
        Tableau avec 5 colonnes et 4 lignes :
        - Ligne 0 : [vide] | Rédaction | Révision | Révision | Approbation
        - Ligne 1 : Nom | [vide]  | Edward NANDA | Armel NGATCHUI | Blaise NDANGANG
        - Ligne 2 : Fonction | [vide] | POGR | RSSI | DSI
        - Ligne 3 : Date & Signature | [vide] | [vide] | [vide] | [vide]
        """
        # Ajouter un espace avant le tableau
        document.add_paragraph()
        
        # Créer le tableau 4 lignes x 5 colonnes
        table = document.add_table(rows=4, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        table.rows[0].cells[0].text = ''
        # Rédaction
        cell_redaction = table.rows[0].cells[1]
        p = cell_redaction.paragraphs[0]
        run = p.add_run('Rédaction')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell_redaction._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        # Révision (fusionner colonnes 2 et 3)
        cell_revision = table.rows[0].cells[2]
        cell_revision_next = table.rows[0].cells[3]
        cell_revision.merge(cell_revision_next)
        p = cell_revision.paragraphs[0]
        run = p.add_run('Révision')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell_revision._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:val'), 'center'
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        # Approbation
        cell_approbation = table.rows[0].cells[4]
        p = cell_approbation.paragraphs[0]
        run = p.add_run('Approbation')
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell_approbation._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        
        # --- Ligne 1 : Noms ---
        # Label
        nom_label = table.rows[1].cells[0]
        nom_label.text = 'Nom'
        if nom_label.paragraphs and nom_label.paragraphs[0].runs:
            nom_label.paragraphs[0].runs[0].bold = True
            nom_label.paragraphs[0].runs[0].font.name = 'Arial'
            nom_label.paragraphs[0].runs[0].font.size = Pt(8)
        tc = nom_label._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        # Noms - Première cellule (Rédaction) vide pour remplissage manuel
        noms = ['', 'Edward NANDA', 'Armel NGATCHUI', 'Blaise NDANGANG']
        for i, nom in enumerate(noms, start=1):
            cell = table.rows[1].cells[i]
            cell.text = nom
            para = cell.paragraphs[0]
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), 'center')
            tcPr.append(tcVAlign)
        
        # --- Ligne 2 : Fonctions ---
        # Label
        fonction_label = table.rows[2].cells[0]
        fonction_label.text = 'Fonction'
        if fonction_label.paragraphs and fonction_label.paragraphs[0].runs:
            fonction_label.paragraphs[0].runs[0].bold = True
            fonction_label.paragraphs[0].runs[0].font.name = 'Arial'
            fonction_label.paragraphs[0].runs[0].font.size = Pt(8)
        tc = fonction_label._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        # Fonctions - Première cellule (Rédaction) vide pour remplissage manuel
        fonctions = ['', 'POGR', 'RSSI', 'DSI']
        for i, fonction in enumerate(fonctions, start=1):
            cell = table.rows[2].cells[i]
            cell.text = fonction
            para = cell.paragraphs[0]
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), 'center')
            tcPr.append(tcVAlign)
        
        # --- Ligne 3 : Date & Signature ---
        # Label
        date_label = table.rows[3].cells[0]
        date_label.text = 'Date & Signature'
        if date_label.paragraphs and date_label.paragraphs[0].runs:
            date_label.paragraphs[0].runs[0].bold = True
            date_label.paragraphs[0].runs[0].font.name = 'Arial'
            date_label.paragraphs[0].runs[0].font.size = Pt(8)
        tc = date_label._element
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:vAlign')
        tcVAlign.set(qn('w:val'), 'center')
        tcPr.append(tcVAlign)
        # Toutes les cellules Date & Signature restent vides pour remplissage manuel
        for i in range(1, 5):
            cell = table.rows[3].cells[i]
            cell.text = '\n\n\n'  # Espace pour la signature
            para = cell.paragraphs[0]
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), 'center')
            tcPr.append(tcVAlign)
    
    def _add_footer_with_page_numbers(self, document):
        """
        Ajoute un bas de page avec un tableau à 3 colonnes et 2 lignes :
        - Colonne 1 : 2 cellules séparées (Version 1.4 / RESTREINT) avec bordures
        - Colonne 2 : Cellule fusionnée verticalement avec texte rouge
        - Colonne 3 : Cellule fusionnée verticalement avec numérotation
        """
        section = document.sections[0]
        footer = section.footer
        
        # Créer un tableau à 3 colonnes et 2 lignes
        footer_table = footer.add_table(rows=2, cols=3, width=Inches(6.5))
        footer_table.autofit = False
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # --- Configurer les propriétés du tableau pour largeur fixe ---
        tbl = footer_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Centrer le tableau
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)
        
        # Définir la largeur totale du tableau (somme des largeurs préférées)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '10340')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Layout fixe pour empêcher l'auto-ajustement
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Définition de la grille des colonnes
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            gridCols = tblGrid.findall(qn('w:gridCol'))
            if len(gridCols) >= 3:
                gridCols[0].set(qn('w:w'), '2193')  # 3,87 cm
                gridCols[1].set(qn('w:w'), '5954')  # 10,5 cm
                gridCols[2].set(qn('w:w'), '2193')  # 3,87 cm
        
        # Respect des dimensions sur Word pour le footer
        col_widths_footer = [2193, 5954, 2193]
        for idx, cell in enumerate(footer_table.rows[0].cells[:3]):
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths_footer[idx]))
            tcW.set(qn('w:type'), 'dxa')
            # Supprimer l'ancien tcW s'il existe
            old_tcW = tcPr.find(qn('w:tcW'))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcPr.append(tcW)
        
        # Aussi pour la ligne 1
        for idx, cell in enumerate(footer_table.rows[1].cells[:3]):
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths_footer[idx]))
            tcW.set(qn('w:type'), 'dxa')
            old_tcW = tcPr.find(qn('w:tcW'))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcPr.append(tcW)
        
        # --- COLONNE 1 : Version 1.4 et RESTREINT (2 cellules séparées) ---
        version_cell = footer_table.rows[0].cells[0]
        version_para = version_cell.paragraphs[0]
        version_para.paragraph_format.space_before = Pt(4)
        version_para.paragraph_format.space_after = Pt(4)
        version_para.paragraph_format.line_spacing = 1.15
        version_run = version_para.add_run('Version 1.4')
        version_run.font.name = 'Arial'
        version_run.font.size = Pt(7)
        version_run.font.bold = True
        version_run.font.color.rgb = RGBColor(0, 0, 0)  # Noir
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc_v = version_cell._element
        tcPr_v = tc_v.get_or_add_tcPr()
        tcVAlign_v = OxmlElement('w:vAlign')
        tcVAlign_v.set(qn('w:val'), 'center')
        tcPr_v.append(tcVAlign_v)
        
        restreint_cell = footer_table.rows[1].cells[0]
        restreint_para = restreint_cell.paragraphs[0]
        restreint_para.paragraph_format.space_before = Pt(4)
        restreint_para.paragraph_format.space_after = Pt(4)
        restreint_para.paragraph_format.line_spacing = 1.15
        restreint_run = restreint_para.add_run('RESTREINT')
        restreint_run.font.name = 'Arial'
        restreint_run.font.size = Pt(7)
        restreint_run.font.bold = True
        restreint_run.font.color.rgb = RGBColor(0, 0, 0)  # Noir
        restreint_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc_r = restreint_cell._element
        tcPr_r = tc_r.get_or_add_tcPr()
        tcVAlign_r = OxmlElement('w:vAlign')
        tcVAlign_r.set(qn('w:val'), 'center')
        tcPr_r.append(tcVAlign_r)
        
        # --- COLONNE 2 : Fusionner les 2 cellules verticalement ---
        center_cell_top = footer_table.rows[0].cells[1]
        center_cell_bottom = footer_table.rows[1].cells[1]
        center_cell_top.merge(center_cell_bottom)
        
        center_para = center_cell_top.paragraphs[0]
        center_para.paragraph_format.space_before = Pt(4)
        center_para.paragraph_format.space_after = Pt(4)
        center_para.paragraph_format.line_spacing = 1.15
        center_run = center_para.add_run('Ce document est la propriété de Boissons du Cameroun')
        center_run.font.name = 'Arial'
        center_run.font.size = Pt(7)
        center_run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge
        center_run.font.bold = False
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Centrer verticalement
        tc_center = center_cell_top._element
        tcPr_center = tc_center.get_or_add_tcPr()
        tcVAlign_center = OxmlElement('w:vAlign')
        tcVAlign_center.set(qn('w:val'), 'center')
        tcPr_center.append(tcVAlign_center)
        
        # --- COLONNE 3 : Fusionner les 2 cellules verticalement ---
        right_cell_top = footer_table.rows[0].cells[2]
        right_cell_bottom = footer_table.rows[1].cells[2]
        right_cell_top.merge(right_cell_bottom)
        
        right_para = right_cell_top.paragraphs[0]
        right_para.paragraph_format.space_before = Pt(4)
        right_para.paragraph_format.space_after = Pt(4)
        right_para.paragraph_format.line_spacing = 1.15
        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter "P a g e "
        page_text_run = right_para.add_run('P a g e ')
        page_text_run.font.name = 'Arial'
        page_text_run.font.size = Pt(7)
        page_text_run.font.bold = True
        page_text_run.font.color.rgb = RGBColor(128, 128, 128)  # Gris
        
        # Numéro de page actuel (en rouge)
        page_num_run = right_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        page_num_run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        page_num_run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_num_run._r.append(fldChar2)
        page_num_run.font.name = 'Arial'
        page_num_run.font.size = Pt(7)
        page_num_run.font.bold = True
        page_num_run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge
        
        # Ajouter " sur "
        sur_run = right_para.add_run(' sur ')
        sur_run.font.name = 'Arial'
        sur_run.font.size = Pt(7)
        sur_run.font.bold = True
        sur_run.font.color.rgb = RGBColor(128, 128, 128)  # Gris
        
        # Nombre total de pages (en rouge)
        total_pages_run = right_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        total_pages_run._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        total_pages_run._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        total_pages_run._r.append(fldChar4)
        total_pages_run.font.name = 'Arial'
        total_pages_run.font.size = Pt(7)
        total_pages_run.font.bold = True
        total_pages_run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge
        
        # Centrer verticalement
        tc_right = right_cell_top._element
        tcPr_right = tc_right.get_or_add_tcPr()
        tcVAlign_right = OxmlElement('w:vAlign')
        tcVAlign_right.set(qn('w:val'), 'center')
        tcPr_right.append(tcVAlign_right)
        
        # --- Configurer les bordures du tableau ---
        tbl = footer_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Bordures du tableau principal
        tblBorders = OxmlElement('w:tblBorders')
        
        # Seules les bordures de la colonne 1 sont visibles
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Bordure fine
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Noir
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Supprimer les bordures des cellules des colonnes 2 et 3
        for row in footer_table.rows:
            for i, cell in enumerate(row.cells):
                if i > 0:  # Colonnes 2 et 3
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        border.set(qn('w:sz'), '0')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'auto')
                        tcBorders.append(border)
                    tcPr.append(tcBorders)

    def _generate_chart_image(self, chart_config, summary_stats):
        """
        Génère une image de graphique à partir de chart_config et summary_stats.
        Retourne le chemin du fichier temporaire créé.
        """
        chart_type = chart_config.get('type')
        title = chart_config.get('title', 'Graphique')
        
        # Créer une nouvelle figure
        fig, ax = plt.subplots(figsize=(8, 5))
        
        try:
            if chart_type == 'bar':
                keys = chart_config.get('keys', [])
                colors = chart_config.get('colors', [])
                orientation = chart_config.get('orientation', 'vertical')
                
                values = [summary_stats.get(key, 0) for key in keys]
                # Convertir les pourcentages en float si nécessaire
                cleaned_values = []
                for v in values:
                    if isinstance(v, str) and '%' in v:
                        cleaned_values.append(float(v.replace('%', '')))
                    else:
                        cleaned_values.append(float(v) if v else 0)
                
                if orientation == 'horizontal':
                    ax.barh(keys, cleaned_values, color=colors if colors else None)
                    ax.set_xlabel('Valeur')
                else:
                    ax.bar(keys, cleaned_values, color=colors if colors else None)
                    ax.set_ylabel('Valeur')
                    plt.xticks(rotation=45, ha='right')
                
                ax.set_title(title, fontsize=14, fontweight='bold')
                plt.tight_layout()
                
            elif chart_type == 'pie':
                keys = chart_config.get('keys', [])
                colors = chart_config.get('colors', [])
                
                values = [summary_stats.get(key, 0) for key in keys]
                # Convertir les pourcentages en float si nécessaire
                cleaned_values = []
                for v in values:
                    if isinstance(v, str) and '%' in v:
                        cleaned_values.append(float(v.replace('%', '')))
                    else:
                        cleaned_values.append(float(v) if v else 0)
                
                ax.pie(cleaned_values, labels=keys, autopct='%1.1f%%', 
                       colors=colors if colors else None, startangle=90)
                ax.set_title(title, fontsize=14, fontweight='bold')
                ax.axis('equal')
                
            elif chart_type == 'gauge':
                key = chart_config.get('key')
                colors = chart_config.get('colors', ['#F44336', '#FF9800', '#4CAF50'])
                thresholds = chart_config.get('thresholds', [50, 80])
                max_value = chart_config.get('max_value', 100)
                
                value = summary_stats.get(key, 0)
                # Extraire le nombre du pourcentage si nécessaire
                if isinstance(value, str):
                    value = float(re.sub(r'[^0-9.]', '', value))
                else:
                    value = float(value)
                
                # Déterminer la couleur selon les seuils
                if value < thresholds[0]:
                    color = colors[0]
                elif value < thresholds[1]:
                    color = colors[1]
                else:
                    color = colors[2]
                
                # Créer un graphique en donut
                remainder = max_value - value
                values = [value, remainder]
                colors_chart = [color, '#E0E0E0']
                
                wedges, texts = ax.pie(values, colors=colors_chart, startangle=90,
                                      wedgeprops=dict(width=0.3))
                
                # Ajouter le texte au centre
                ax.text(0, 0, f'{value:.1f}%', ha='center', va='center',
                       fontsize=24, fontweight='bold')
                ax.set_title(title, fontsize=14, fontweight='bold')
                ax.axis('equal')
            
            # Sauvegarder dans un fichier temporaire
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            plt.savefig(temp_file.name, dpi=150, bbox_inches='tight')
            plt.close(fig)
            
            return temp_file.name
            
        except Exception as e:
            plt.close(fig)
            print(f"Erreur lors de la génération du graphique '{title}': {e}")
            return None
    
    def generate_and_save_report(self, user_data, control_data, analysis_results, save_path, period_label='N/A', control_metadata=None, execution_date=None):
        """
        Génère un rapport DOCX programmatiquement en utilisant python-docx.
        Inclut les graphiques générés avec matplotlib.
        
        Args:
            control_metadata: Métadonnées du contrôle définies dans le script
            execution_date: Date d'exécution au format YYYYMMDD-HHMMSS
        """
        document = docx.Document()
        temp_chart_files = []  # Pour nettoyer les fichiers temporaires

        try:
            # --- Définir le style de base du document ---
            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            # Contenue du document en Arial 8pt noir
            content_style = document.styles.add_style('Content8pt', WD_STYLE_TYPE.PARAGRAPH)
            content_style.font.name = 'Arial'
            content_style.font.size = Pt(8)
            content_style.font.color.rgb = RGBColor(0, 0, 0)

            # Titres en Arial 8pt, gras, souligné, noir
            for level in range(1, 4):
                heading_style = document.styles[f'Heading {level}']
                heading_style.font.name = 'Arial'
                heading_style.font.size = Pt(8)
                heading_style.font.color.rgb = RGBColor(0, 0, 0)
                if level == 1:
                    heading_style.font.bold = True
                    heading_style.font.underline = True

            # --- Ajouter l'en-tête avec logo et tableau ---
            control_name = control_data.get('name', 'N/A')
            control_code = control_data.get('code', 'CTL_SSI_02_SAVE_2025_10_3')
            self._add_header_with_logo_and_table(
                document, 
                control_name, 
                control_code, 
                analysis_results,
                control_metadata=control_metadata,
                execution_date=execution_date
            )
            
            # --- Ajouter le bas de page ---
            self._add_footer_with_page_numbers(document)

            # --- NOUVELLE STRUCTURE DU CORPS DU RAPPORT ---
            
            # 1. Description du contrôle
            document.add_heading('Description du contrôle :', level=1)
            # Appliquer Arial 8pt, gras, souligné, noir
            heading_para = document.paragraphs[-1]
            for run in heading_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            description_text = "N/A"
            if control_metadata and 'description' in control_metadata:
                description_text = control_metadata['description']
            p_desc = document.add_paragraph(description_text)
            p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p_desc.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 2. Analyse
            document.add_heading('Analyse :', level=1)
            # Appliquer Arial 8pt, gras, souligné, noir
            heading_para = document.paragraphs[-1]
            for run in heading_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            analyse_text = "N/A"
            if control_metadata and 'analyse' in control_metadata:
                analyse_text = control_metadata['analyse']
            p_analyse = document.add_paragraph(analyse_text)
            p_analyse.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p_analyse.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 3. Résultats
            document.add_heading('Résultats :', level=1)
            # Appliquer Arial 8pt, gras, souligné, noir
            heading_para = document.paragraphs[-1]
            for run in heading_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # --- Corps du rapport (boucle sur les résultats d'analyse) ---
            if not analysis_results:
                 document.add_paragraph("L'analyse n'a produit aucun resultat a afficher.")
            
            for section in analysis_results:
                # Titre de la section (niveau 2 car sous "Résultats")
                document.add_heading(section.get('title', 'Section de resultat'), level=2)

                # Statistiques résumées
                summary_stats = section.get('summary_stats', {})
                if summary_stats:
                    document.add_heading('Statistiques Cles', level=3)
                for key, value in summary_stats.items():
                    p_stat = document.add_paragraph(style='List Bullet')
                    p_stat.add_run(f"{key}: ").bold = True
                    p_stat.add_run(str(value))
                    for run in p_stat.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)                # Graphiques (si présents)
                chart_configs = section.get('chart_configs', [])
                if chart_configs and summary_stats:
                    document.add_heading('Graphiques', level=3)
                    
                    for chart_config in chart_configs:
                        chart_path = self._generate_chart_image(chart_config, summary_stats)
                        if chart_path:
                            temp_chart_files.append(chart_path)
                            try:
                                # Ajouter l'image
                                document.add_picture(chart_path, width=Inches(2.75))
                                
                                # Centrer l'image
                                last_paragraph = document.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                            except Exception as e:
                                print(f"Erreur lors de l'ajout du graphique au document: {e}")
                
                # Tableau des données détaillées
                items = section.get('items', [])
                display_columns = section.get('display_columns', {})
                
                if items and display_columns:
                    document.add_heading('Donnees Detaillees', level=3)
                    
                    # Création du tableau
                    if isinstance(display_columns, dict):
                        headers = list(display_columns.values())
                        column_keys = list(display_columns.keys())
                    else:
                        # Si display_columns est une liste de dicts avec 'key' et 'label'
                        headers = [col.get('label', col.get('key', '')) for col in display_columns]
                        column_keys = [col.get('key', '') for col in display_columns]
                    
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    table.autofit = False
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Configurer largeur fixe comme le footer
                    tbl = table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'center')
                    tblPr.append(jc)
                    
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '10340')
                    tblW.set(qn('w:type'), 'dxa')
                    tblPr.append(tblW)
                    
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'fixed')
                    tblPr.append(tblLayout)

                    # Remplissage des en-têtes
                    header_cells = table.rows[0].cells
                    for i, header_text in enumerate(headers):
                        cell_paragraph = header_cells[i].paragraphs[0]
                        run = cell_paragraph.add_run(str(header_text))
                        run.bold = True
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(6)
                        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Remplissage des données (limiter à 60 lignes pour éviter un rapport trop lourd)
                    max_rows = min(60, len(items))
                    for item in items[:max_rows]:
                        row_cells = table.add_row().cells
                        for i, key in enumerate(column_keys):
                            row_cells[i].text = str(item.get(key, ''))
                            para = row_cells[i].paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.name = 'Tahoma'
                                run.font.size = Pt(6)
                    
                    if len(items) > max_rows:
                        p_note = document.add_paragraph()
                        p_note.add_run(f"Note : Seules les {max_rows} premieres lignes sont affichees. ").italic = True
                        p_note.add_run(f"Total de {len(items)} lignes dans le fichier Excel complet.").italic = True
                        for run in p_note.runs:
                            run.font.size = Pt(8)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                
            
            # 4. Recommandations (vide - à remplir manuellement)
            document.add_heading('Recommandations :', level=1)
            # Appliquer Arial 8pt, gras, souligné, noir
            heading_para = document.paragraphs[-1]
            for run in heading_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            document.add_paragraph()  # Espace vide pour remplissage manuel
            p = document.paragraphs[-1]
            p.style = document.styles['Content8pt']

            
            # 5. Évidence de suivi des exceptions (vide - à remplir manuellement)
            document.add_heading('Évidence de suivi des exceptions :', level=1)
            # Appliquer Arial 8pt, gras, souligné, noir
            heading_para = document.paragraphs[-1]
            for run in heading_para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            document.add_paragraph()  # Espace vide pour remplissage manuel
            p = document.paragraphs[-1]
            p.style = document.styles['Content8pt']

            
            # 6. Tableau de signatures
            self._add_signature_table(document)

            # --- Sauvegarde du document final ---
            document.save(save_path)
            
            return save_path
            
        except Exception as e:
            raise IOError(f"Impossible de sauvegarder le fichier de rapport : {e}")
        
        finally:
            # Nettoyer les fichiers temporaires de graphiques
            for temp_file in temp_chart_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except Exception as e:
                    print(f"Erreur lors de la suppression du fichier temporaire {temp_file}: {e}")

report_service = ReportGenerator()